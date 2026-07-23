# gerar_documentacao.py
"""Gera a documentação viva do projeto em Word (.docx) na raiz do projeto.
Rodar com:  python gerar_documentacao.py
Atualize o conteúdo abaixo conforme o projeto evoluir e rode de novo —
o arquivo é sobrescrito, preservando o histórico na seção 'Atualizações'."""

from __future__ import annotations

import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ARQUIVO_SAIDA = "../Documentacao_Sistema_CBSM.docx"

AZUL_CBSM = RGBColor(0x0F, 0x17, 0x2A)
CINZA = RGBColor(0x47, 0x55, 0x69)


def _titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL_CBSM
    return h


def _paragrafo(doc, texto, negrito=False, italico=False, tamanho=11, cor=None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    run.font.size = Pt(tamanho)
    if cor:
        run.font.color.rgb = cor
    return p


def _bullet(doc, texto, negrito_prefixo=None):
    p = doc.add_paragraph(style="List Bullet")
    if negrito_prefixo:
        r = p.add_run(negrito_prefixo)
        r.bold = True
    p.add_run(texto)
    return p


def _tabela(doc, cabecalho, linhas):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Light Grid Accent 1"
    for i, titulo_col in enumerate(cabecalho):
        cel = t.rows[0].cells[i]
        cel.text = titulo_col
        for p in cel.paragraphs:
            for r in p.runs:
                r.bold = True
    for linha in linhas:
        row = t.add_row()
        for i, valor in enumerate(linha):
            row.cells[i].text = str(valor)
    doc.add_paragraph("")
    return t


def gerar():
    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    # Capa
    titulo = doc.add_heading("Sistema CBSM — Passivo & Receita Diferida", level=0)
    for run in titulo.runs:
        run.font.color.rgb = AZUL_CBSM
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run("Conciliação de pontos Dotz (IFRS 15 / CPC 47) — Documentação Viva do Projeto")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = CINZA
    _paragrafo(doc, f"Última atualização: {datetime.date.today().strftime('%d/%m/%Y')}",
               italico=True, tamanho=10, cor=CINZA)
    doc.add_paragraph("")

    # 1. Visão geral
    _titulo(doc, "1. O que o sistema faz")
    _paragrafo(doc, "Concilia o passivo e a receita diferida de pontos de fidelidade (Dotz), conforme "
                     "IFRS 15 / CPC 47, cruzando duas fontes:")
    _bullet(doc, "balancete do SAP (aba Trial Balance) — lado contábil.", "Lado contábil: ")
    _bullet(doc, "planilha interna U1 - Movimentação — lado controle, com passivo recalculado por conta.",
            "Lado controle: ")
    _paragrafo(doc, "A partir dessa conciliação, o sistema gera lançamentos contábeis de ajuste para "
                     "cobrir os desvios encontrados, e oferece um dashboard Streamlit navegável por abas "
                     "para explorar cada parte da planilha original.")

    # 2. Motores de receita
    _titulo(doc, "2. Motores de receita")
    _paragrafo(doc, "Conciliação (Controle U1 / balancete) — 4 motores:", negrito=True)
    _tabela(doc, ["Motor", "O que é", "Tipo de diferimento"], [
        ["Breakage", "Expiração estimada de pontos que se tornam receita.",
         "Diferimento 48 meses (linear por competência de emissão)"],
        ["Spread", "Margem sobre a venda de pontos.",
         "Diferimento 48 meses (linear por competência de emissão)"],
        ["Trocas (Custo do Produto)", "Resgate de pontos + provisão manual \"a faturar\" (conta 3111199).",
         "Fato gerador (resgate)"],
        ["Promodotz", "Pontos promocionais.", "Fato gerador (evento)"],
    ])
    _paragrafo(doc, "Receita reconhecida no resultado (gráficos \"Receita por Categoria\") — base "
                     "DIFERENTE, por decisão do usuário: sempre que o app mostra receita no resultado, "
                     "a fonte é a coluna \"Classificação Receita\" da aba DRE 2026 "
                     "(data_loader.load_dre_classificacao_receita), não a Base Receita Motor. São "
                     "5 categorias:", negrito=True)
    _tabela(doc, ["Categoria", "Origem na DRE", "Tipo de diferimento"], [
        ["Breakage", "Classificação Receita = Breakage", "Diferimento 48 meses"],
        ["Spread", "Classificação Receita = Spread", "Diferimento 48 meses"],
        ["Trocas", "Classificação Receita = Receita de Trocas (exceto conta 3111160)", "Fato gerador (resgate)"],
        ["Receita Projetos Especiais", "Conta 3111160 (\"Receita de Resgate de Dotz - Split Fee e "
         "Projetos Especiais\"), separada do restante de Receita de Trocas", "Fato gerador (resgate)"],
        ["Promodotz", "EXCLUÍDO destes gráficos — é redutor de custo, não receita "
         "(conta 3211010, \"(-) Custo de Troca Promodotz\")", "—"],
    ])
    _paragrafo(doc, "Atenção: o \"Trocas\" da DRE (gráficos de resultado) é um valor DIFERENTE do "
                     "\"Trocas (Custo do Produto)\" da conciliação/Controle U1 — vêm de bases contábeis "
                     "distintas. O app deixa isso explícito na aba \"Receita Diferida por Categoria\".",
               italico=True, cor=CINZA)

    # 3. Arquitetura
    _titulo(doc, "3. Arquitetura")
    _paragrafo(doc, "Camadas do sistema (todas em data_raw/):", negrito=True)
    _tabela(doc, ["Arquivo", "Camada", "Responsabilidade"], [
        ["config.py", "Config", "Caminho do Excel (dados.xlsm desde 20/07/2026 — base passou a ser "
         "atualizada por macros; antes era dados.xlsx) e nomes das abas usadas."],
        ["data_loader.py", "Ingestão (RAW)", "Lê balancete, Base Receita Motor, U1 e as abas de apoio "
         "(U1.4, U1.5, U1.6, Dados_Gráfico, Fat. Análise de Vendas). Tipagem robusta, sem alterar valores."],
        ["data_pipeline.py", "Negócio", "Agrega receita, monta a DRE, concilia por blocos "
         "(Quadros A/G), calcula passivo contábil × controle. Não altera o RAW."],
        ["gerador_lancamentos.py", "Negócio", "Gera lançamentos de ajuste (partida dobrada: "
         "Débito=J se J>0, Crédito=-J se J<0) para os 4 motores (Breakage, Custo do Produto, "
         "Spread, Promodotz). Fonte do valor: o \"check\" de cada grupo de negócio da aba "
         "Controle U1 (Passivo Total × Passivo Recalculado), pelo valor exato, sem zerar "
         "diferenças pequenas — existe para que essa aba sempre feche, reconhecendo a receita "
         "correspondente. Descrições de conta lidas do balancete (Trial Balance), não digitadas."],
        ["app.py", "Apresentação", "Dashboard Streamlit interativo — 9 abas (ver seção 5). Inclui "
         "gate de upload do dados.xlsx/dados.xlsm (_garantir_arquivo_excel) para funcionar também "
         "quando hospedado sem o arquivo local (Streamlit Community Cloud). Toda tabela do app tem "
         "um botão \"📥 Exportar (Excel)\" (_botao_exportar_excel/to_excel_bytes) que gera .xlsx "
         "formatado (cabeçalho em negrito, coluna congelada, largura ajustada) — não depende do "
         "ícone de CSV nativo do st.dataframe, que continua existindo ao lado (limitação do "
         "próprio Streamlit, não removível pelo código)."],
        ["export_html.py", "Apresentação (opcional)", "Snapshot HTML estático AUTOCONTIDO (logo e "
         "Plotly.js embutidos em base64, funciona offline) espelhando as 9 abas do app.py, com as "
         "mesmas cores Dotz e a mesma base de receita (DRE). Só quando pedido explicitamente — o "
         "app Streamlit ao vivo continua sendo o alvo real; o HTML não tem filtros interativos."],
        ["gerar_documentacao.py", "Documentação", "Gera este documento. Rodar de novo após "
         "mudanças relevantes no projeto."],
        ["gerar_manual_operacional.py", "Documentação", "Gera o Manual Operacional (.docx) — "
         "passo a passo de uso para quem opera o sistema no dia a dia (atualizar dados, reiniciar "
         "o app, gerar o HTML, pedir melhorias). Documento separado da documentação técnica."],
        ["gerar_logo_placeholder.py", "Identidade visual (opcional)", "Gera um ícone Dotz "
         "placeholder (círculo laranja + \"DZ\") via Pillow, enquanto o arquivo oficial da "
         "marca não é exportado — ver seção 6."],
    ])
    _paragrafo(doc, "Arquitetura híbrida: os números ATUAIS/de conciliação vêm do balancete e da U1 "
                     "(via data_loader). A PROJEÇÃO usa um motor atuarial de safras "
                     "(apropriação linear em 48 meses), com dados-fonte nas abas U1.4/U1.5.")

    # 4. Regras contábeis
    _titulo(doc, "4. Regras contábeis importantes")
    _bullet(doc, "No balancete SAP, passivo (grupo 2) e receita (grupo 3) vêm com crédito NEGATIVO. "
                  "O RAW preserva o sinal do SAP; a inversão para \"sentido de negócio\" (positivo) só "
                  "ocorre na camada de negócio (data_pipeline.py).")
    _bullet(doc, "A Base Receita Motor já vem positiva (sentido de negócio) — nunca inverter o sinal dela.")
    _bullet(doc, "Contas de passivo: grupos 219 (circulante) e 231 (não circulante). Usar só contas "
                  "ANALÍTICAS de 7 dígitos, para evitar dupla contagem pai/filho da hierarquia SAP.")
    _bullet(doc, "Split CP/LP não é derivável por prefixo de conta. Fórmula validada via leitura "
                  "das FÓRMULAS reais do Excel (openpyxl, data_only=False): U1 - Movimentação!G60 = "
                  "G8+G10+G19+G37 (idem I60 para LP) — soma dos passivo_cp/passivo_lp dos 4 "
                  "cabeçalhos de bloco (Subtotal Split Fee, Breakage, Custo do Produto, Spread). "
                  "Implementado em data_loader.calcular_passivo_cp_lp(), bate ao centavo com a "
                  "mesma fonte que a aba Dados_Gráfico usa (coluna \"CBSM\"), mas SEM os ajustes de "
                  "Netpoints/Dotz Pay de lá (fora do escopo deste sistema) — eliminou a dependência "
                  "da aba Dados_Gráfico para este número e corrigiu um bug pré-existente do "
                  "dashboard (que antes lia a coluna \"total_ajustado\", incluindo Netpoints/Dotz "
                  "Pay indevidamente).")
    _bullet(doc, "_to_float tolera formato BR (1.234,56) e US/Excel (1234.56) — decide pela presença "
                  "de vírgula. Nunca remover o ponto decimal cegamente (já inflou valores 100x no passado).")
    _bullet(doc, "Contas contábeis são sempre TEXTO, nunca número (ex.: '3111160', não 3111160.0).")
    _bullet(doc, "Premissas contábeis (régua de breakage, CPD, provisão manual) são inputs explícitos "
                  "no Excel — nunca adivinhados pelo código.")
    _bullet(doc, "Tolerâncias de conciliação são em CAMADAS, por decisão explícita do usuário "
                  "(20/07/2026): o Passivo Diferido Total (quadros A/G em data_pipeline.conciliar) "
                  "tem que fechar exato, tolerância R$ 1,00 (TOLERANCIA). Já as linhas por conta/"
                  "motor (\"Receita/Breakage\", \"Receita/Custo do Produto\" etc.) e o check "
                  "conta-a-conta da própria U1 toleram até R$ 999,00 (TOLERANCIA_CONTA) — ruído de "
                  "arredondamento não deve travar a conciliação geral. Já o check por GRUPO de "
                  "negócio na aba Controle U1 (app.EPS_CHECK_LINHA) usa uma terceira tolerância, "
                  "R$ 10,00, independente das duas de cima.")

    # 5. Abas do dashboard
    _titulo(doc, "5. As 9 abas do dashboard (app.py)")
    abas = [
        ("📊 Visão Geral / Conciliação", "Alerta verde/vermelho, cards de passivo (contábil/controle/delta), "
         "gráfico de receita por categoria (base DRE), rosca de Composição do Passivo Diferido, "
         "tabela de Conciliação (Total Passivo / Passivo Recalculado / Receita por bloco) e, "
         "incorporado ao final da aba, o conteúdo da antiga aba Dados_Gráfico: quadro "
         "\"Expectativa de Realização da Receita\" (vintage × ano de realização) e os quadros "
         "Consolidado / Receita Diferida CP/LP."),
        ("📉 Passivo Diferido", "Composição do passivo (rosca + tabela por bloco) e o split "
         "Circulante (CP) × Não Circulante (LP) — hoje calculado direto da fórmula da U1 - "
         "Movimentação (ver seção 4), sem depender da aba Dados_Gráfico — conferido contra o "
         "Passivo Total do Controle U1."),
        ("💰 Receita Diferida por Categoria", "Base DRE 2026 (5 categorias, ver seção 2). Filtros "
         "de competência e categoria, gráfico de barras empilhadas com linha de Receita Bruta "
         "total, resumo por competência × categoria, e detalhe expansível linha a linha com o "
         "tipo de diferimento de cada categoria."),
        ("🧾 Faturamento de Pontos", "Resumo do faturamento do ano por categoria (valor de vendas "
         "e quantidade de pontos), com detalhe por competência × categoria recolhido e exportável "
         "em Excel."),
        ("🔧 Controle U1", "Conciliação por grupo de negócio: 4 cartões (nomeados com o texto "
         "literal da planilha) cruzando Passivo Total × Passivo Recalculado × Check × Receita, "
         "cada um expansível a nível de conta contábil."),
        ("U1.4_Par._Emissão Expiração", "Motor atuarial do Breakage: quadro \"Receita a Reconhecer "
         "por Safra × Competência\" (2022-2026, com total) + gráfico de barras empilhadas, check "
         "de conciliação com a DRE (mês a mês e acumulado, considerando o mês de apuração "
         "corrente), e a série completa por safra em ordem decrescente (2026 primeiro)."),
        ("U1.5_Par._Emissão Margem", "Mesma estrutura da U1.4, para o motor de Spread."),
        ("U1.6_Emissão_Resgates", "Série mensal de resgate de pontos e o quadro de conferência "
         "\"Tie-in Contábil × Analyser\"."),
        ("🧾 Lançamentos de Ajuste", "Diário de ajuste gerado a partir dos desvios de conciliação, "
         "com checks de partida dobrada e cobertura, e exportação em Excel."),
    ]
    for nome, desc in abas:
        _bullet(doc, desc, f"{nome}: ")

    # 6. Identidade visual
    _titulo(doc, "6. Identidade visual (Design System Dotz)")
    _paragrafo(doc, "Paleta oficial da marca (fornecida pelo usuário), aplicada em app.py:", negrito=True)
    _tabela(doc, ["Cor", "Hex", "Uso no dashboard"], [
        ["Laranja (primária)", "#FF4F0D", "Breakage, status OK-adjacente em gráficos, DRE, "
         "rampa de safra/vintage, categoria NSF do faturamento."],
        ["Amarelo", "#FEC114", "Trocas (Custo do Produto), categoria SF do faturamento."],
        ["Verde", "#009F3C", "Spread, status \"OK\" da tabela de Conciliação, categoria BV."],
        ["Preto", "#000000", "Promodotz, categoria MKT do faturamento."],
        ["Magenta", "#D82598", "Categoria Viaja Dotz do faturamento."],
        ["Vermelho escuro", "#AF1010", "Status \"DIVERGENTE\", categoria Comissão com "
         "Conversão de Dotz."],
        ["Bege", "#E3D2C8", "Categoria Prestação de Serviços Comissão - Banco do Brasil."],
    ])
    _paragrafo(doc, "Composição do Passivo Diferido (rosca, abas Visão Geral e Passivo Diferido): "
                     "Promodotz permanece na rosca (não é removido) e recebe destaque visual "
                     "(fatia levemente separada — \"pull\") a pedido do usuário, mesmo sabendo que "
                     "Breakage + Custo do Produto + Spread já somam 100% do Passivo Diferido Total "
                     "sozinhos.")
    _paragrafo(doc, "Logo: o header (app.py) já está preparado para exibir logo_dotz.png / "
                     "icone_dotz.png (pasta data_raw) via st.logo + st.image, com fallback "
                     "silencioso caso os arquivos não existam. Como o arquivo oficial ainda não "
                     "foi exportado, gerar_logo_placeholder.py cria uma aproximação (círculo "
                     "laranja #FF4F0D + \"DZ\" em negrito) — trocar pelo arquivo oficial assim "
                     "que disponível, para fidelidade total à tipografia da marca.")

    # 7. Números validados
    _titulo(doc, "7. Números de referência validados (jun/2026 YTD)")
    _tabela(doc, ["Item", "Valor"], [
        ["Passivo Diferido Total", "R$ 196.321.706,82"],
        ["Receita Breakage", "R$ 21.610.402,09"],
        ["Receita Custo do Produto / Trocas", "R$ 22.665.859,03"],
        ["Receita Spread", "R$ 16.095.728,24"],
        ["Receita Promodotz", "R$ 1.075.793,82"],
        ["Passivo Circulante (CP)", "R$ 135.027.553,55"],
        ["Passivo Não Circulante (LP)", "R$ 61.294.153,27"],
        ["Tolerância — Passivo Total (quadros A/G)", "R$ 1,00"],
        ["Tolerância — por conta/motor e check U1", "R$ 999,00 (ver seção 4)"],
        ["Tolerância — check por grupo (aba Controle U1)", "R$ 10,00"],
    ])

    # 8. Ambiente
    _titulo(doc, "8. Ambiente técnico")
    _bullet(doc, "A pasta do projeto está no OneDrive com caminho longo, o que estoura o limite de "
                  "~260 caracteres do Windows e quebra a instalação do Streamlit num .venv local.")
    _bullet(doc, "Solução: ambiente virtual em caminho curto — %USERPROFILE%\\venvs\\sistema_cbsm "
                  "— com pandas, openpyxl, streamlit, plotly, python-docx e Pillow instalados.")
    _bullet(doc, "Rodar o dashboard: abrir PowerShell na pasta data_raw e executar "
                  "\"%USERPROFILE%\\venvs\\sistema_cbsm\\Scripts\\streamlit.exe run app.py\".")
    _bullet(doc, "O arquivo de dados deve se chamar dados.xlsm (desde 20/07/2026 — antes era "
                  "dados.xlsx) e estar na mesma pasta dos .py (data_raw). Fechar o Excel antes de "
                  "eu (ou o app) precisar ler o arquivo — o Python recebe \"Permission denied\" "
                  "se o arquivo estiver aberto.")
    _bullet(doc, "Cuidado ao reabrir um terminal novo: o login do Windows na máquina pode ser um "
                  "usuário diferente do dono da pasta OneDrive (ex.: login local \"sd.client\" "
                  "com OneDrive logado como \"william.abreu_dotz\") — nesse caso "
                  "$env:USERPROFILE no PowerShell resolve para a pasta errada e o caminho do "
                  ".venv (%USERPROFILE%\\venvs\\sistema_cbsm) não é encontrado. Usar o caminho "
                  "completo (C:\\Users\\william.abreu_dotz\\venvs\\sistema_cbsm\\...) evita o problema.")

    # 9. Publicação e acesso
    _titulo(doc, "9. Publicação e acesso")
    _paragrafo(doc, "GitHub (backup do código-fonte):", negrito=True)
    _bullet(doc, "Repositório: github.com/williamabreu-pixel/Receita-Diferida-CBSM, branch main. "
                  "Público (decisão consciente: nenhum dado financeiro real é versionado — "
                  "dados.xlsx, dashboard.html e os .docx gerados estão no .gitignore desde o "
                  "primeiro commit).")
    _bullet(doc, "requirements.txt e README.md criados para permitir que qualquer colaborador "
                  "clone o repositório e rode localmente (precisa fornecer seu próprio dados.xlsx).")
    _paragrafo(doc, "Cópia em rede (Google Drive compartilhado):", negrito=True)
    _bullet(doc, "A pasta local do projeto (dentro do OneDrive) é sincronizada automaticamente por "
                  "um Google Drive for Desktop configurado na máquina, aparecendo também em "
                  "G:\\Drives compartilhados\\Controladoria\\...\\2.1 - Receita Diferida\\"
                  "Sistema_CBSM — inclusive o dados.xlsx real, aceitável porque é uma pasta de "
                  "acesso restrito da empresa (diferente do GitHub). Não é necessário copiar nada "
                  "manualmente para lá.")
    _paragrafo(doc, "Acesso local em rede (Streamlit rodando na máquina):", negrito=True)
    _bullet(doc, "Regra de Firewall do Windows criada (\"Streamlit CBSM 8501\", TCP 8501, entrada, "
                  "todos os perfis) permitindo acesso de outras máquinas na mesma rede local via "
                  "http://<IP-da-máquina>:8501.")
    _bullet(doc, "Limitação importante: só funciona para quem estiver na MESMA rede física/Wi-Fi "
                  "da máquina que roda o app — um IP de rede local (ex.: 192.168.x.x) não é "
                  "alcançável de fora dessa rede (nem por VPN, nem de outro escritório). Também "
                  "depende da máquina ficar ligada com o processo do Streamlit rodando.")
    _bullet(doc, "Testado em 20/07/2026 na rede Wi-Fi corporativa (\"DOTZ 6\", IP 172.31.10.87) — "
                  "MESMO ASSIM outras máquinas não conseguiram conectar (ERR_CONNECTION_TIMED_OUT), "
                  "apesar do firewall liberado e do processo escutando em todas as interfaces. "
                  "Causa provável: isolamento de cliente (AP/Client Isolation) no Wi-Fi corporativo "
                  "— uma proteção de rede que a TI normalmente não desativa, e que está fora do "
                  "controle deste projeto. Reforça o Streamlit Community Cloud como o caminho "
                  "correto para acesso de verdade pelo time.")
    _paragrafo(doc, "Streamlit Community Cloud (URL pública, independe de rede/máquina ligada):",
               negrito=True)
    _bullet(doc, "app.py recebeu um gate de upload (_garantir_arquivo_excel): se dados.xlsx não "
                  "existir no disco do servidor, a tela pede upload manual do arquivo, que fica só "
                  "na sessão do navegador (nunca é persistido).")
    _bullet(doc, "Status em 16/07/2026: deploy tentado em share.streamlit.io, mas bloqueado. "
                  "Causa raiz identificada via DevTools (aba Network do navegador): a chamada do "
                  "backend do Streamlit para o GitHub (api/v2/github/query-repository) retorna "
                  "404 — o GitHub App do Streamlit nunca foi instalado de fato na conta "
                  "(distinto de uma autorização OAuth comum, que foi concedida mas não resolve "
                  "isso). Repositório já foi tornado público para eliminar a hipótese de "
                  "permissão de repo privado, sem sucesso. Próximo passo: instalar o GitHub App "
                  "do Streamlit diretamente pela Marketplace do GitHub (github.com/apps/...), "
                  "escolhendo \"All repositories\" — retomar isso na próxima sessão.")

    # 10. Itens em aberto
    _titulo(doc, "10. Itens em aberto")
    _bullet(doc, "Concluir a publicação no Streamlit Community Cloud — falta instalar o GitHub App "
                  "corretamente (ver seção 9); é o passo que falta para ter uma URL pública "
                  "funcionando, independente de rede local ou da máquina ligada.")
    _bullet(doc, "Segundo tipo de lançamento de ajuste (reclassificação CP × LP) em "
                  "gerador_lancamentos.py — os 4 motores de receita já têm lançamento completo "
                  "(20/07/2026), mas a reclassificação entre Circulante e Não Circulante ainda "
                  "não tem lançamento próprio; falta decidir se/como gerar essa entrada.")
    _bullet(doc, "Substituir o logo placeholder (gerar_logo_placeholder.py) pelo arquivo oficial "
                  "da marca (logo_dotz.png / icone_dotz.png) assim que exportado — hoje é só uma "
                  "aproximação, sem a tipografia customizada do \"Z\" do logo real.")
    _bullet(doc, "Limpar os scripts de investigação usados para descobrir a fórmula do CP/LP "
                  "(inspecionar_*.py, testar_cp_lp.py, validar_cp_lp_real.py, *_dump.txt) — já "
                  "excluídos do Git via .gitignore, mas ainda presentes no disco/pasta de rede.")

    # 11. Histórico de atualizações
    _titulo(doc, "11. Histórico de atualizações")
    _paragrafo(doc, "Documentação viva — cada sessão de trabalho relevante deve somar uma entrada "
                     "aqui (nova entrada no topo, com data fixa) antes de regerar o arquivo. As datas "
                     "abaixo são fixas no texto — não usar datetime.date.today() para elas, senão "
                     "o histórico troca de data toda vez que o documento é regerado.",
               italico=True, cor=CINZA)
    _bullet(doc, "to_excel_bytes (app.py) passou a formatar todo .xlsx exportado (cabeçalho em "
                  "negrito com fundo escuro, coluna congelada, largura ajustada ao conteúdo). "
                  "Novo helper _botao_exportar_excel adicionado a TODAS as tabelas do dashboard "
                  "que ainda só tinham o ícone de CSV nativo do st.dataframe (Conciliação, DRE, "
                  "Receita por Categoria — resumo e detalhe —, Passivo Diferido, Faturamento por "
                  "categoria, Controle U1 por grupo, série atuarial e check DRE de U1.4/U1.5, "
                  "grades brutas, U1.6 e Dados_Gráfico) — 48 botões de exportação no total. "
                  "Validado de ponta a ponta via streamlit.testing.v1.AppTest (roda o app sem "
                  "navegador): 0 exceções nas 9 abas. O ícone de CSV nativo continua aparecendo "
                  "ao lado — é do próprio Streamlit, não há como removê-lo pelo código.",
              "22/07/2026 — ")
    _bullet(doc, "Base de dados migrada para dados.xlsm (a planilha passou a ser atualizada por "
                  "macros/Power Query, alimentada pelos exports brutos do SAP em "
                  "data_raw/sap_entradas/) — .gitignore atualizado para *.xlsm e sap_entradas/ "
                  "antes de qualquer outra coisa (as 39 abas são idênticas às do dados.xlsx "
                  "anterior, nenhuma mudança estrutural necessária além do nome do arquivo). "
                  "Tolerâncias de conciliação reorganizadas em camadas (R$ 1 para o Passivo "
                  "Total, R$ 999 para linhas por conta/motor e check da U1, R$ 10 para o check "
                  "por grupo na aba Controle U1). gerador_lancamentos.py reescrito: cobre os 4 "
                  "motores (antes só Breakage e Promodotz), corrigido um bug de partida dobrada "
                  "(as 2 linhas de cada lançamento caíam sempre do mesmo lado, débito nunca "
                  "fechava com crédito), fonte do valor passou a ser o check por grupo da aba "
                  "Controle U1 (não mais o pequeno desvio da conciliação), sempre pelo valor "
                  "exato, e as descrições de conta foram conferidas contra o balancete real. "
                  "Coluna J_desvio removida da exibição/exportação da aba Lançamentos. Testada "
                  "publicação em rede Wi-Fi corporativa (DOTZ 6) — bloqueada por isolamento de "
                  "cliente do Wi-Fi, fora do controle deste projeto; reforça o Streamlit "
                  "Community Cloud como caminho definitivo (ainda pendente).", "20/07/2026 — ")
    _bullet(doc, "Split CP/LP corrigido com a fórmula real da U1 - Movimentação (lida via "
                  "openpyxl), eliminando a dependência da aba Dados_Gráfico e corrigindo um bug "
                  "pré-existente do dashboard (Netpoints/Dotz Pay entrando indevidamente no "
                  "número). Receita por Categoria migrada para a base DRE 2026 (coluna "
                  "Classificação Receita), com a nova categoria \"Receita Projetos Especiais\" "
                  "separada (conta 3111160) e Promodotz excluído dos gráficos de resultado. "
                  "Projeto publicado no GitHub (repositório público, sem dados reais). "
                  "export_html.py reescrito para espelhar as 9 abas completas do app com as cores "
                  "Dotz. Manual Operacional criado (gerar_manual_operacional.py). Tentativas de "
                  "publicação: rede local (funciona só na mesma rede física) e Streamlit Community "
                  "Cloud (bloqueado por instalação pendente do GitHub App — ver seção 9).",
              "16/07/2026 — ")
    _bullet(doc, "Identidade visual Dotz aplicada em app.py com a paleta OFICIAL da marca "
                  "(#FF4F0D laranja, #FEC114 amarelo, #009F3C verde, #000000 preto, #D82598 "
                  "magenta, #AF1010 vermelho, #E3D2C8 bege) — motores, roscas (Promodotz mantido "
                  "e destacado, não removido), DRE, status da Conciliação, rampa de safra/vintage "
                  "(U1.4/U1.5/Dados_Gráfico) e as 7 categorias do Faturamento de Pontos. Header "
                  "com logo Dotz preparado (fallback silencioso) e ícone placeholder gerado via "
                  "Pillow enquanto o arquivo oficial não é exportado.", "16/07/2026 — ")
    _bullet(doc, "Documento inicial gerado, consolidando todo o estado do projeto até aqui: "
                  "10 abas do dashboard, split CP/LP resolvido, aba Controle U1 redesenhada em "
                  "4 grupos de negócio, aba Faturamento de Pontos criada, exportações padronizadas "
                  "em Excel (XLSX).", "15/07/2026 — ")

    doc.save(ARQUIVO_SAIDA)
    print(f"Documento gerado: {ARQUIVO_SAIDA}")


if __name__ == "__main__":
    gerar()
