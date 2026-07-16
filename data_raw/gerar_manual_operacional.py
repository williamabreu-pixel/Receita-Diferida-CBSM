# gerar_manual_operacional.py
"""Gera o Manual Operacional (.docx) na raiz do projeto — passo a passo de
como atualizar dados mensais e solicitar melhorias no app, para quem opera
o sistema no dia a dia (não é documentação técnica, é o manual de uso).
Rodar com:  python gerar_manual_operacional.py
Atualize o conteúdo abaixo conforme o processo mudar e rode de novo — o
arquivo é sobrescrito."""

from __future__ import annotations

import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ARQUIVO_SAIDA = "../Manual_Operacional_Sistema_CBSM.docx"

LARANJA_DOTZ = RGBColor(0xFF, 0x4F, 0x0D)
AZUL_CBSM = RGBColor(0x0F, 0x17, 0x2A)
CINZA = RGBColor(0x47, 0x55, 0x69)

CAMINHO_PROJETO = (r"C:\Users\william.abreu_dotz\OneDrive - CBSM - Companhia Brasileira de "
                   r"Solucoes de Marketing\Área de Trabalho\Sistema_CBSM")
CAMINHO_DADOS = CAMINHO_PROJETO + r"\data_raw\dados.xlsx"
CAMINHO_VENV = r"C:\Users\william.abreu_dotz\venvs\sistema_cbsm"
URL_LOCAL = "http://localhost:8501"
URL_REDE = "http://192.168.0.122:8501"


def _titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = AZUL_CBSM
    return h


def _paragrafo(doc, texto, negrito=False, italico=False, tamanho=11, cor=None):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    run.font.size = Pt(tamanho)
    if cor:
        run.font.color.rgb = cor
    return p


def _bullet(doc, texto, negrito_prefixo=None):
    p = doc.add_paragraph(style="List Bullet")
    if negrito_prefixo:
        r = p.add_run(negrito_prefixo)
        r.bold = True
    p.add_run(texto)
    return p


def _passo(doc, texto, negrito_prefixo=None):
    p = doc.add_paragraph(style="List Number")
    if negrito_prefixo:
        r = p.add_run(negrito_prefixo)
        r.bold = True
    p.add_run(texto)
    return p


def _codigo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(texto)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return p


def _caixa_atencao(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    run = p.add_run(f"⚠️ {texto}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xAF, 0x10, 0x10)
    return p


def _tabela(doc, cabecalho, linhas):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Light Grid Accent 1"
    for i, titulo_col in enumerate(cabecalho):
        cel = t.rows[0].cells[i]
        cel.text = titulo_col
        for p in cel.paragraphs:
            for r in p.runs:
                r.bold = True
    for linha in linhas:
        row = t.add_row()
        for i, valor in enumerate(linha):
            row.cells[i].text = str(valor)
    return t


def montar_documento():
    doc = Document()

    titulo = doc.add_heading("Manual Operacional — Sistema CBSM", level=0)
    for run in titulo.runs:
        run.font.color.rgb = LARANJA_DOTZ
    _paragrafo(doc, "Passivo & Receita Diferida — IFRS 15 / CPC 47", tamanho=14, cor=CINZA, italico=True)
    _paragrafo(doc, f"Atualizado em: {datetime.date.today().strftime('%d/%m/%Y')}", tamanho=9, cor=CINZA)
    _paragrafo(doc, "Este manual é para quem opera o sistema no dia a dia (sem necessidade de "
                    "conhecimento técnico) — cobre como atualizar os dados mensais e como pedir "
                    "ajustes ou melhorias no app dashboard.")

    # 1. Visão geral -----------------------------------------------------
    _titulo(doc, "1. Visão Geral")
    _paragrafo(doc, "O dashboard roda como um app Streamlit nesta própria máquina. Dois endereços "
                    "dão acesso a ele:")
    _bullet(doc, URL_LOCAL, negrito_prefixo="Nesta máquina: ")
    _bullet(doc, f"{URL_REDE}  (funciona só para quem está na mesma rede Wi-Fi/local)",
           negrito_prefixo="Colegas na rede: ")
    _caixa_atencao(doc, "O app só fica acessível enquanto esta máquina estiver ligada e o processo "
                        "do Streamlit estiver rodando. Se o computador reiniciar ou hibernar, o "
                        "acesso cai até alguém subir o app de novo (seção 4).")

    # 2. Atualização de dados ---------------------------------------------
    _titulo(doc, "2. Atualizando os Dados Mensais")
    _paragrafo(doc, "Sempre que houver um novo fechamento (novo mês, ajuste na planilha, etc.), "
                    "siga estes passos, nesta ordem:")
    _passo(doc, "Substitua o arquivo pela nova versão, MANTENDO o mesmo nome (dados.xlsx) e "
                "MANTENDO todas as abas originais (mesmos nomes de aba).",
          negrito_prefixo="Salve o novo Excel exatamente neste caminho: ")
    _codigo(doc, CAMINHO_DADOS)
    _passo(doc, "Abra o dashboard no navegador (endereço da seção 1).")
    _passo(doc, "Clique no menu ☰ no canto superior direito da página.")
    _passo(doc, "Clique em \"Clear cache\".")
    _passo(doc, "Clique em \"Rerun\" (ou aperte a tecla R).")
    _caixa_atencao(doc, "Este passo de \"Clear cache\" + \"Rerun\" é obrigatório. O Streamlit guarda "
                        "os resultados calculados em memória — só trocar o arquivo dados.xlsx NÃO "
                        "atualiza os números na tela sozinho. Sem isso, todo mundo continua vendo "
                        "os valores antigos mesmo com o arquivo já trocado.")
    _passo(doc, "Confira, na aba \"Visão Geral / Conciliação\", se o status voltou "
                "\"🟢 Conciliação 100% OK\" e se os valores fazem sentido para o novo mês.")

    # 3. Pasta de rede ------------------------------------------------------
    _titulo(doc, "3. Cópia na Rede (Google Drive Compartilhado)")
    _paragrafo(doc, "A pasta do projeto também aparece automaticamente em:")
    _codigo(doc, r"G:\Drives compartilhados\Controladoria\Controladoria\05 - Contabilidade\3 - 2026"
                 r"\1 - Fechamento\01. CBSM\06. Junho\02 - Passivo\2.1 - Receita Diferida\Sistema_CBSM")
    _paragrafo(doc, "Essa cópia é só um espelho automático (sincronização do Google Drive) — não "
                    "precisa copiar nada manualmente para lá. Ela serve como backup e para "
                    "colaboradores consultarem os arquivos-fonte, mas não é essa cópia que o "
                    "dashboard ao vivo usa (ele lê da pasta local da seção 2).")

    # 4. Reiniciar o app ------------------------------------------------------
    _titulo(doc, "4. Reiniciando o App (se ele parar de responder)")
    _paragrafo(doc, "Se os links da seção 1 pararem de abrir, o processo do Streamlit não está mais "
                    "rodando. Para subir de novo:")
    _passo(doc, "Abra o PowerShell.")
    _passo(doc, "Cole e rode o comando abaixo (entra na pasta do projeto):")
    _codigo(doc, f'cd "{CAMINHO_PROJETO}\\data_raw"')
    _passo(doc, "Cole e rode o comando abaixo (sobe o app, acessível na rede):")
    _codigo(doc, f'& "{CAMINHO_VENV}\\Scripts\\python.exe" -m streamlit run app.py '
                 f'--server.address=0.0.0.0 --server.port 8501 --server.headless true')
    _passo(doc, "Deixe essa janela do PowerShell aberta — fechá-la derruba o app. Espere a "
                "mensagem confirmando que o Streamlit subiu e teste o link.")

    # 5. Instantâneo HTML -----------------------------------------------------
    _titulo(doc, "5. Gerando um Instantâneo em HTML (para enviar por e-mail)")
    _paragrafo(doc, "Se precisar mandar o dashboard para alguém que não está na mesma rede (por "
                    "e-mail, por exemplo), gere um arquivo único que abre em qualquer navegador, "
                    "sem precisar do app rodando:")
    _passo(doc, "Abra o PowerShell e rode:")
    _codigo(doc, f'cd "{CAMINHO_PROJETO}\\data_raw"')
    _codigo(doc, f'& "{CAMINHO_VENV}\\Scripts\\python.exe" export_html.py')
    _passo(doc, "O arquivo é gerado em: " + CAMINHO_PROJETO + r"\dashboard.html")
    _caixa_atencao(doc, "Esse arquivo é um retrato fixo (sem os filtros interativos do app) e "
                        "contém os números reais da CBSM — trate-o com o mesmo cuidado do "
                        "dados.xlsx: não é para subir no GitHub nem em lugar público.")

    # 6. Pedindo melhorias -----------------------------------------------------
    _titulo(doc, "6. Pedindo Ajustes ou Melhorias no App")
    _paragrafo(doc, "Você não precisa saber programar — só descrever o que precisa. Para o pedido "
                    "ser atendido rápido e sem retrabalho, inclua:")
    _bullet(doc, "Em qual aba do dashboard é (ex.: \"Receita Diferida por Categoria\").", negrito_prefixo="1. ")
    _bullet(doc, "O que está errado OU o que você gostaria que existisse (ex.: \"o gráfico deveria "
                "separar X de Y\", \"faltou uma coluna com Z\").", negrito_prefixo="2. ")
    _bullet(doc, "Se possível, um print de tela mostrando o problema ou o exemplo do Excel que "
                "mostra como deveria ficar.", negrito_prefixo="3. ")
    _bullet(doc, "Qual a régra de negócio por trás (ex.: \"Promodotz não entra como receita, é "
                "redutor de custo\") — isso evita que o ajuste saia errado.", negrito_prefixo="4. ")
    _paragrafo(doc, "Depois de qualquer alteração no código do app, sempre teste antes de considerar "
                    "concluído: recarregue o dashboard (Clear cache + Rerun) e confira a aba afetada.")

    # 7. Onde fica cada coisa ---------------------------------------------------
    _titulo(doc, "7. Onde Fica Cada Coisa")
    _tabela(doc, ["Item", "Onde"],
           [
               ["Planilha de dados (dados.xlsx)", r"...\Sistema_CBSM\data_raw\dados.xlsx (local)"],
               ["App do dashboard (código)", r"...\Sistema_CBSM\data_raw\app.py"],
               ["Cópia de backup na rede", r"G:\...\2.1 - Receita Diferida\Sistema_CBSM (espelho automático)"],
               ["Backup do código-fonte (sem dados)", "GitHub privado — williamabreu-pixel/Receita-Diferida-CBSM"],
               ["Este manual", r"...\Sistema_CBSM\Manual_Operacional_Sistema_CBSM.docx"],
               ["Documentação técnica completa", r"...\Sistema_CBSM\Documentacao_Sistema_CBSM.docx"],
           ])

    # 8. Segurança ------------------------------------------------------------
    _titulo(doc, "8. Observações de Segurança")
    _bullet(doc, "O dados.xlsx (e o dashboard.html gerado a partir dele) contém números financeiros "
               "reais da CBSM — nunca enviar para o GitHub, nem para ferramentas públicas na internet.")
    _bullet(doc, "O link de rede (http://192.168.0.122:8501) só funciona para quem está na mesma "
               "rede Wi-Fi/local — não é acessível pela internet.")
    _bullet(doc, "Qualquer mudança de firewall ou de rede depende da TI — não tente liberar portas "
               "sozinho sem alinhar com eles.")

    doc.save(ARQUIVO_SAIDA)
    print(f"Manual gerado: {ARQUIVO_SAIDA}")


if __name__ == "__main__":
    montar_documento()
